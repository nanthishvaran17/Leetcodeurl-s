"""
Master Word Report Generator (.docx)
Generates official landscape Word performance reports for Nandha Engineering College.
Consumes ONLY the canonical dataset dictionary.
"""
import io
import datetime
from typing import Optional, Dict, Any, List
from sqlalchemy.orm import Session

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from backend.config.report_config import BATCH_CONFIG, DEPARTMENT_COORDINATORS, get_coordinator_for_department


def set_cell_background(cell, fill_hex: str):
    """Utility to set XML background color for docx table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_table_borders(table, color="94A3B8", sz="4", val="single"):
    """Applies clean borders to all cells in a docx table."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), val)
                border.set(qn('w:sz'), sz)
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), color)
                tcBorders.append(border)
            tcPr.append(tcBorders)


def build_weekly_performance_docx(data: Dict[str, Any], dept_id: Optional[int] = None) -> bytes:
    """
    Builds official landscape Word performance report directly from the canonical dataset.
    Does NOT query database.
    Includes all four batches: 2026-2030, 2025-2029, 2024-2028, 2023-2027.
    """
    if not DOCX_AVAILABLE:
        buffer = io.BytesIO()
        content = f"NANDHA ENGINEERING COLLEGE, ERODE - 638 052.\n" \
                  f"Leetcode Performance - Weekly Report\n" \
                  f"Date: {data.get('report_date', datetime.date.today().strftime('%d.%m.%Y'))}\n\n"
        buffer.write(content.encode('utf-8'))
        buffer.seek(0)
        return buffer.getvalue()

    doc = Document()

    # Configure Landscape A4 layout
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    report_date_str = data.get("report_date", datetime.date.today().strftime("%d.%m.%Y"))
    dept_summaries = data.get("dept_summaries", [])

    if dept_id is not None:
        dept_summaries = [d for d in dept_summaries if d.get("department_id") == dept_id]

    if not dept_summaries:
        dept_summaries = [{
            "department_id": 1,
            "department": "CSE(CS)",
            "department_name": "Department of Computer Science and Engineering (Cyber Security)",
            "coordinator": DEPARTMENT_COORDINATORS.get("CSE(CS)"),
            "batches": {}
        }]

    for dept_index, dept in enumerate(dept_summaries):
        if dept_index > 0:
            doc.add_page_break()

        dept_code = dept.get("department", "CSE")
        dept_title = dept.get("department_name") or f"Department of {dept_code}"
        if not dept_title.startswith("Department of"):
            dept_title = f"Department of {dept_title}"
        coordinator = dept.get("coordinator") or get_coordinator_for_department(dept_code)

        # ── 1. OFFICIAL INSTITUTIONAL HEADER ──
        p_inst = doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_inst.paragraph_format.space_after = Pt(2)
        r_inst = p_inst.add_run("NANDHA ENGINEERING COLLEGE, ERODE - 638 052.")
        r_inst.font.name = "Times New Roman"
        r_inst.font.size = Pt(13)
        r_inst.font.bold = True
        r_inst.font.color.rgb = RGBColor(15, 23, 42)

        p_aff = doc.add_paragraph()
        p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_aff.paragraph_format.space_after = Pt(2)
        r_aff = p_aff.add_run("(An Autonomous Institution, Affiliated to Anna University, Chennai)")
        r_aff.font.name = "Times New Roman"
        r_aff.font.size = Pt(9.5)
        r_aff.font.italic = True
        r_aff.font.color.rgb = RGBColor(71, 85, 105)

        p_dept = doc.add_paragraph()
        p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_dept.paragraph_format.space_after = Pt(2)
        r_dept = p_dept.add_run(dept_title)
        r_dept.font.name = "Times New Roman"
        r_dept.font.size = Pt(11)
        r_dept.font.bold = True
        r_dept.font.color.rgb = RGBColor(30, 41, 59)

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.paragraph_format.space_after = Pt(2)
        r_date = p_date.add_run(f"Date: {report_date_str}")
        r_date.font.name = "Times New Roman"
        r_date.font.size = Pt(9.5)
        r_date.font.bold = True

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_title.paragraph_format.space_after = Pt(2)
        r_title = p_title.add_run("Leetcode Performance - Weekly Report")
        r_title.font.name = "Times New Roman"
        r_title.font.size = Pt(10)
        r_title.font.bold = True
        r_title.font.underline = True

        p_coord = doc.add_paragraph()
        p_coord.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_coord.paragraph_format.space_after = Pt(6)
        r_coord = p_coord.add_run(f"Name & Designation of the Academic Coordinator: {coordinator}")
        r_coord.font.name = "Times New Roman"
        r_coord.font.size = Pt(9.5)
        r_coord.font.bold = True

        # ── 2. OFFICIAL MULTI-LEVEL TABLE ──
        batches_dict = dept.get("batches", {})
        active_batch_cfgs = [b for b in BATCH_CONFIG if batches_dict.get(b["key"], {}).get("total_students", 0) > 0]
        if not active_batch_cfgs:
            active_batch_cfgs = BATCH_CONFIG  # Fallback

        total_table_rows = 2 + (len(active_batch_cfgs) * 2)
        table = doc.add_table(rows=total_table_rows, cols=13)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        col_widths = [
            Inches(1.2), Inches(0.85),
            Inches(0.68), Inches(0.68), Inches(0.68), Inches(0.68), Inches(0.65),
            Inches(0.62), Inches(0.62), Inches(0.62), Inches(0.62),
            Inches(0.9), Inches(0.9)
        ]

        for row in table.rows:
            for c_idx, width in enumerate(col_widths):
                row.cells[c_idx].width = width

        # Row 0 (Main Headers)
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 0).paragraphs[0].text = "Batch"

        table.cell(0, 1).merge(table.cell(1, 1))
        table.cell(0, 1).paragraphs[0].text = "No. of Students\n(Total Count)"

        table.cell(0, 2).merge(table.cell(0, 6))
        table.cell(0, 2).paragraphs[0].text = "Number of Problems Solved"

        table.cell(0, 7).merge(table.cell(0, 10))
        table.cell(0, 7).paragraphs[0].text = "Weekly Contest Attended"

        table.cell(0, 11).merge(table.cell(0, 12))
        table.cell(0, 11).paragraphs[0].text = "Leetcode Contest Rating and Ranking"

        # Row 1 (Sub-Headers)
        sub_headers = [
            (2, "Above 500"),
            (3, "250 - 500"),
            (4, "100 - 249"),
            (5, "1 - 99"),
            (6, "0"),
            (7, "4Q"),
            (8, "3Q"),
            (9, "2Q"),
            (10, "1Q"),
            (11, "Rating:\n> 1500"),
            (12, "Ranking:\n< 20000")
        ]

        for col_idx, sub_title in sub_headers:
            table.cell(1, col_idx).paragraphs[0].text = sub_title

        # Style Header Rows
        for r_idx in range(2):
            for c_idx in range(13):
                cell = table.cell(r_idx, c_idx)
                set_cell_background(cell, "0F172A")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(8.5)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

        # ── 3. DATA ROWS FOR EACH BATCH (LAST WEEK & CURRENT WEEK) ──
        curr_row = 2

        for b_cfg in active_batch_cfgs:
            b_key = b_cfg["key"]
            b_label = b_cfg["label"]
            b_metrics = batches_dict.get(b_key, {})
            lw = b_metrics.get("last_week", {})
            cw = b_metrics.get("current_week", {})
            tot_st = b_metrics.get("total_students", 0)

            # Last Week Row
            last_week_cells = table.rows[curr_row].cells
            last_week_vals = [
                f"{b_label}\n(Last Week)",
                str(tot_st),
                str(lw.get("prob_above_500", 0)),
                str(lw.get("prob_250_500", 0)),
                str(lw.get("prob_100_249", 0)),
                str(lw.get("prob_1_99", 0)),
                str(lw.get("prob_0", 0)),
                str(lw.get("q4", 0)),
                str(lw.get("q3", 0)),
                str(lw.get("q2", 0)),
                str(lw.get("q1", 0)),
                str(lw.get("rating_above_1500", 0)),
                str(lw.get("rank_below_20000", 0))
            ]

            # Current Week Row
            curr_week_cells = table.rows[curr_row + 1].cells
            curr_week_vals = [
                f"{b_label}\n(Current Week)",
                str(tot_st),
                str(cw.get("prob_above_500", 0)),
                str(cw.get("prob_250_500", 0)),
                str(cw.get("prob_100_249", 0)),
                str(cw.get("prob_1_99", 0)),
                str(cw.get("prob_0", 0)),
                str(cw.get("q4", 0)),
                str(cw.get("q3", 0)),
                str(cw.get("q2", 0)),
                str(cw.get("q1", 0)),
                str(cw.get("rating_above_1500", 0)),
                str(cw.get("rank_below_20000", 0))
            ]

            for c_i in range(13):
                # Last week cell
                p_l = last_week_cells[c_i].paragraphs[0]
                p_l.text = last_week_vals[c_i]
                p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last_week_cells[c_i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for run in p_l.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)
                    if c_i in (0, 1):
                        run.font.bold = True

                # Current week cell
                p_c = curr_week_cells[c_i].paragraphs[0]
                p_c.text = curr_week_vals[c_i]
                p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                curr_week_cells[c_i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_background(curr_week_cells[c_i], "F8FAFC")
                for run in p_c.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)
                    if c_i in (0, 1):
                        run.font.bold = True

            curr_row += 2

        set_table_borders(table, color="94A3B8", sz="4", val="single")

        # ── 4. SIGNATURE BLOCK ──
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(14)

        p_sig = doc.add_paragraph()
        p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_sig.paragraph_format.space_after = Pt(2)
        r_sig = p_sig.add_run(
            "Verified Signatures:       "
            "Academic Coordinator                                                                        "
            "Head of Department"
        )
        r_sig.font.name = "Times New Roman"
        r_sig.font.size = Pt(9.5)
        r_sig.font.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_word_report(db: Session, dept_id: Optional[int] = None, *args, **kwargs) -> bytes:
    """
    Compatibility wrapper: produces Word document from canonical dataset.
    """
    from backend.services.weekly_report_service import generate_weekly_performance_data
    current_user = kwargs.get('current_user')
    data = generate_weekly_performance_data(db, current_user=current_user)
    return build_weekly_performance_docx(data, dept_id=dept_id)


generate_snapshot_word_report = generate_word_report
generate_weekly_word_report = generate_word_report
