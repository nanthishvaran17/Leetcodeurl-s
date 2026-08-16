import os
import io
import datetime
import re
from typing import Optional

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def set_cell_background(cell, fill_hex: str):
    """Utility to set XML background color for docx table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_document_margins(doc, top=0.75, bottom=0.75, left=0.75, right=0.75):
    """Utility to set margins for all sections in docx."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def export_word_from_dataset(dataset: dict) -> bytes:
    """
    WORD DOCX EXPORTER
    Generates editable Microsoft Word (.docx) document directly from normalized ReportDataset.
    Uses Times New Roman font ONLY.
    """
    if not DOCX_AVAILABLE:
        buffer = io.BytesIO()
        buffer.write(f"Report: {dataset.get('title', 'Universal Report')}\n".encode('utf-8'))
        buffer.seek(0)
        return buffer.getvalue()

    doc = Document()
    set_document_margins(doc, 0.75, 0.75, 0.75, 0.75)

    # 1. Header Banner with Logo
    logo_path = os.path.join(os.path.dirname(__file__), "..", "assets", "nandha_emblem.png")
    if os.path.exists(logo_path):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Inches(1.1))
        except Exception:
            pass

    p_college = doc.add_paragraph()
    p_college.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_coll = p_college.add_run("NANDHA ENGINEERING COLLEGE (AUTONOMOUS)")
    r_coll.font.name = "Times New Roman"
    r_coll.font.size = Pt(16)
    r_coll.font.bold = True
    r_coll.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Approved by AICTE, New Delhi & Affiliated to Anna University, Chennai | Erode - 638 052, Tamil Nadu")
    r_sub.font.name = "Times New Roman"
    r_sub.font.size = Pt(9)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(2, 132, 199)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(dataset.get('title', 'Universal Performance Report'))
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(13)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run(f"Report ID: {dataset.get('reportId', '')}   |   Generated: {dataset.get('generatedAt', '')[:10]}   |   Status: {dataset.get('dataStatus', 'READY')}")
    r_meta.font.name = "Times New Roman"
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2. Executive Summary Metrics Table
    metrics = dataset.get("metrics", {})
    if metrics:
        doc.add_heading("I. Executive Summary Metrics", level=2)
        p_h = doc.paragraphs[-1]
        for r in p_h.runs:
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(15, 23, 42)

        metric_items = list(metrics.items())
        table_m = doc.add_table(rows=len(metric_items), cols=2)
        table_m.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, (k, v) in enumerate(metric_items):
            row_cells = table_m.rows[r_idx].cells

            p_k = row_cells[0].paragraphs[0]
            label_text = re.sub(r'([A-Z])', r' \1', str(k)).strip().title()
            r_k = p_k.add_run(label_text)
            r_k.font.name = "Times New Roman"
            r_k.font.size = Pt(10)
            r_k.font.bold = True

            p_v = row_cells[1].paragraphs[0]
            p_v.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            val_str = f"{v:,}" if isinstance(v, (int, float)) and v > 999 else str(v if v is not None else "—")
            r_v = p_v.add_run(val_str)
            r_v.font.name = "Times New Roman"
            r_v.font.size = Pt(10)
            r_v.font.bold = True
            r_v.font.color.rgb = RGBColor(2, 132, 199)

            if r_idx % 2 == 0:
                set_cell_background(row_cells[0], "F8FAFC")
                set_cell_background(row_cells[1], "F8FAFC")

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 3. Weekly Contest Participation Matrix (for Weekly Contest datasets)
    contest_rows = dataset.get("rows", [])
    is_weekly = (
        dataset.get("report_type", "").lower() in ("weekly_contest", "weekly contest")
        or bool(contest_rows)
    )

    if is_weekly and contest_rows:
        doc.add_heading("II. Weekly Contest Participation Matrix", level=2)
        p_hw = doc.paragraphs[-1]
        for r in p_hw.runs:
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(15, 23, 42)

        STATUS_LABEL = {
            "PUBLIC_ATTENDED":    "PUBLIC",
            "ATTENDED":           "PUBLIC",
            "VIRTUAL_ATTENDED":   "VIRTUAL",
            "PUBLIC_NOT_ATTENDED":"NOT ATTENDED",
            "NOT_ATTENDED":       "NOT ATTENDED",
            "DATA_ERROR":         "DATA ERROR",
            "PENDING":            "PENDING",
        }

        w_headers = ["S.No", "Reg No", "Student Name", "Dept", "Yr", "Status", "Contest Name", "Q1", "Q2", "Q3", "Q4", "Contest Solved", "Rank"]
        table_w = doc.add_table(rows=1 + len(contest_rows), cols=len(w_headers))
        table_w.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table_w.rows[0].cells
        for i, h in enumerate(w_headers):
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.name = "Times New Roman"
            run.font.size = Pt(8.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(hdr_cells[i], "0F172A")

        for idx, r in enumerate(contest_rows, start=1):
            row_cells = table_w.rows[idx].cells
            p_status = r.get("participation_status", "NOT_ATTENDED")
            attended = p_status in ("PUBLIC", "PUBLIC_ATTENDED", "ATTENDED", "VIRTUAL", "VIRTUAL_ATTENDED")
            status_str = STATUS_LABEL.get(p_status, p_status)
            c_name_val = r.get("contest_name") or dataset.get("contestName") or "Weekly Contest"

            def _parse_q(val):
                if isinstance(val, (int, float)):
                    return 1 if val > 0 else 0
                if str(val).strip().isdigit():
                    return 1 if int(val) > 0 else 0
                return 0

            q1 = _parse_q(r.get("q1"))
            q2 = _parse_q(r.get("q2"))
            q3 = _parse_q(r.get("q3"))
            q4 = _parse_q(r.get("q4"))
            solved_cnt = q1 + q2 + q3 + q4

            r_vals = [
                str(idx),
                r.get("reg_no", ""),
                r.get("name", ""),
                r.get("dept", ""),
                str(r.get("year", "")),
                status_str,
                c_name_val,
                str(q1) if attended else "—",
                str(q2) if attended else "—",
                str(q3) if attended else "—",
                str(q4) if attended else "—",
                str(solved_cnt) if attended else "—",
                str(r.get("rank") or r.get("contest_rank") or "—") if attended else "—",
            ]

            bg_hex = "ECFDF5" if attended else "FFF1F2"
            for i, val in enumerate(r_vals):
                set_cell_background(row_cells[i], bg_hex)
                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i not in (2, 6) else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                run.font.name = "Times New Roman"
                run.font.size = Pt(8)
                if i in (0, 5, 11):
                    run.font.bold = True
                if i == 5:
                    run.font.color.rgb = RGBColor(6, 95, 70) if attended else RGBColor(153, 27, 27)

        doc.add_paragraph().paragraph_format.space_after = Pt(10)



    # 3. Problem Solving Category Summary Table
    distribution = dataset.get("distribution")
    if distribution:
        doc.add_heading("II. Problem Solving Category Summary", level=2)
        p_h2 = doc.paragraphs[-1]
        for r in p_h2.runs:
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(15, 23, 42)

        dist_items = list(distribution.items())
        table_d = doc.add_table(rows=1 + len(dist_items), cols=2)
        table_d.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table_d.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run("Category Range").font.bold = True
        hdr_cells[1].paragraphs[0].add_run("Student Count").font.bold = True
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_background(hdr_cells[0], "0F172A")
        set_cell_background(hdr_cells[1], "0F172A")
        for cell in hdr_cells:
            for r in cell.paragraphs[0].runs:
                r.font.name = "Times New Roman"
                r.font.color.rgb = RGBColor(255, 255, 255)

        for idx, (cat, cnt) in enumerate(dist_items, start=1):
            row_cells = table_d.rows[idx].cells
            row_cells[0].paragraphs[0].add_run(str(cat)).font.name = "Times New Roman"

            pv = row_cells[1].paragraphs[0]
            pv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rv = pv.add_run(str(cnt))
            rv.font.name = "Times New Roman"
            rv.font.bold = True

            if idx % 2 == 1:
                set_cell_background(row_cells[0], "F1F5F9")
                set_cell_background(row_cells[1], "F1F5F9")

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 4. Top Performers Leaderboard Table
    top_students = dataset.get("topStudents")
    if top_students:
        doc.add_heading("III. Top Performers Leaderboard", level=2)
        p_h3 = doc.paragraphs[-1]
        for r in p_h3.runs:
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(15, 23, 42)

        headers = ["Rank", "Reg No", "Student Name", "Dept", "Year", "Easy", "Med", "Hard", "Total Solved", "Rating"]
        table_top = doc.add_table(rows=1 + len(top_students), cols=len(headers))
        table_top.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table_top.rows[0].cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(hdr_cells[i], "0F172A")

        for idx, s in enumerate(top_students, start=1):
            row_cells = table_top.rows[idx].cells
            row_vals = [
                f"#{idx}",
                s.get("reg_no", ""),
                s.get("name", ""),
                s.get("dept", ""),
                s.get("year", ""),
                str(s.get("easy", 0)),
                str(s.get("medium", 0)),
                str(s.get("hard", 0)),
                str(s.get("total_solved", 0)),
                f"{round(s['rating'], 1):,}" if s.get("rating") else "Unrated"
            ]
            for i, val in enumerate(row_vals):
                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 2 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5)
                if i in (0, 8):
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(15, 23, 42)

            if idx % 2 == 1:
                for c in row_cells:
                    set_cell_background(c, "F1F5F9")

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 5. Full Roster Table
    all_students = dataset.get("allStudents")
    if all_students:
        doc.add_heading("IV. Student Performance Master Roster", level=2)
        p_h4 = doc.paragraphs[-1]
        for r in p_h4.runs:
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(15, 23, 42)

        headers = ["S.No", "Reg No", "Student Name", "Dept", "Year", "Total Solved", "Status"]
        table_all = doc.add_table(rows=1 + len(all_students), cols=len(headers))
        table_all.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table_all.rows[0].cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(hdr_cells[i], "0F172A")

        for idx, s in enumerate(all_students, start=1):
            row_cells = table_all.rows[idx].cells
            row_vals = [
                str(idx),
                s.get("reg_no", ""),
                s.get("name", ""),
                s.get("dept", ""),
                s.get("year", ""),
                str(s.get("total_solved") if s.get("total_solved") is not None else "—"),
                s.get("status", "UNVERIFIED")
            ]
            for i, val in enumerate(row_vals):
                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 2 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5)
                if i == 5:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(2, 132, 199)
                if i == 6 and val == "UNVERIFIED":
                    run.font.color.rgb = RGBColor(220, 38, 38)
                elif i == 6 and val == "VERIFIED":
                    run.font.color.rgb = RGBColor(16, 185, 129)

            if idx % 2 == 1:
                for c in row_cells:
                    set_cell_background(c, "F8FAFC")

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 6. Official Contest Participations Table
    participations = dataset.get("participations")
    if participations:
        doc.add_heading("V. Official Contest Participation Log", level=2)
        p_h5 = doc.paragraphs[-1]
        for r in p_h5.runs:
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(15, 23, 42)

        headers = ["S.No", "Contest Name", "Date", "Reg No", "Student Name", "Dept", "Score", "Rank"]
        table_p = doc.add_table(rows=1 + len(participations), cols=len(headers))
        table_p.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table_p.rows[0].cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(hdr_cells[i], "0F172A")

        for idx, p_item in enumerate(participations, start=1):
            row_cells = table_p.rows[idx].cells
            row_vals = [
                str(idx),
                p_item.get("contest_name", ""),
                p_item.get("date", ""),
                p_item.get("reg_no", ""),
                p_item.get("student_name", ""),
                p_item.get("dept", ""),
                f"{p_item.get('problems_solved', 0)} / {p_item.get('total_problems', 4)}",
                str(p_item.get("rank", "-"))
            ]
            for i, val in enumerate(row_vals):
                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i not in (1, 4) else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5)
                if i in (1, 6):
                    run.font.bold = True

            if idx % 2 == 1:
                for c in row_cells:
                    set_cell_background(c, "F8FAFC")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
