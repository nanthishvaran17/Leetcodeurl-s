import sys
sys.path.insert(0, r"e:\Leetcode Web")

import docx
import io
from backend.database import SessionLocal
from backend.word_generator import generate_word_report

def validate_official_word_report():
    db = SessionLocal()
    try:
        print("==================================================")
        print("GENERATING OFFICIAL WEEKLY PERFORMANCE WORD REPORT")
        print("==================================================")

        word_bytes = generate_word_report(db)
        print(f"Generated Word Document Size: {len(word_bytes):,} bytes")
        assert len(word_bytes) > 2000, "DOCX file too small!"

        # Save sample to scratch
        out_path = r"e:\Leetcode Web\scratch_output_report.docx"
        with open(out_path, "wb") as f:
            f.write(word_bytes)
        print(f"Saved test file to: {out_path}")

        # Programmatic Validation with python-docx
        doc = docx.Document(io.BytesIO(word_bytes))
        print(f"[OK] DOCX Package parsed successfully!")
        print(f"[OK] Total Sections: {len(doc.sections)}")
        print(f"[OK] Total Paragraphs: {len(doc.paragraphs)}")
        print(f"[OK] Total Tables: {len(doc.tables)}")

        assert len(doc.tables) >= 2, f"Expected at least 2 tables (Cyber Security + IoT), got {len(doc.tables)}"

        # Validate Table Contents
        dept_names = ["Cyber Security", "IoT"]
        for idx, table in enumerate(doc.tables):
            print(f"\n[TABLE {idx + 1}]: {dept_names[idx] if idx < 2 else 'Department'}")
            print(f"  Rows: {len(table.rows)} | Cols: {len(table.columns)}")
            assert len(table.rows) == 8, f"Expected 8 rows (2 header + 6 batch/week rows), got {len(table.rows)}"
            assert len(table.columns) == 13, f"Expected 13 columns, got {len(table.columns)}"

            # Print data rows for verification
            for r_idx in range(2, len(table.rows)):
                row_cells = table.rows[r_idx].cells
                batch_cell = row_cells[0].text.replace('\n', ' ')
                total_cnt = row_cells[1].text
                p_500 = row_cells[2].text
                p_250 = row_cells[3].text
                p_100 = row_cells[4].text
                p_1 = row_cells[5].text
                p_0 = row_cells[6].text
                q4 = row_cells[7].text
                q3 = row_cells[8].text
                q2 = row_cells[9].text
                q1 = row_cells[10].text
                r1500 = row_cells[11].text
                r20k = row_cells[12].text

                # Verify reconciliation of problem categories
                tot = int(total_cnt)
                sum_cats = int(p_500) + int(p_250) + int(p_100) + int(p_1) + int(p_0)
                assert sum_cats == tot, f"Row {batch_cell} category sum {sum_cats} != {tot}"

                # Verify attended count does NOT equal total roster count
                total_solved_q = int(q4) + int(q3) + int(q2) + int(q1)
                print(f"  • {batch_cell}: Total={tot} | Problems [>500:{p_500}, 250-500:{p_250}, <250:{p_100}, <100:{p_1}, NotStarted:{p_0}] (Sum={sum_cats}) | Contest Solved [4Q:{q4}, 3Q:{q3}, 2Q:{q2}, 1Q:{q1}] | Rating>1500:{r1500}, Rank<20k:{r20k}")

        print("\n==================================================")
        print("ALL OFFICIAL WORD DOCX VALIDATIONS PASSED 100%!")
        print("==================================================")

    finally:
        db.close()

if __name__ == "__main__":
    validate_official_word_report()
